import json
import logging
import fitz
import traceback
from typing import List
from interviewai.db.index import UserIndexHelper, InterviewNamespace, index
from interviewai.firebase import get_fs_client
from datetime import datetime
from io import BytesIO
from google.cloud import storage
from google.cloud import firestore as fstore
from interviewai.config.config import get_config
from langchain.document_loaders.base import BaseLoader
from langchain.docstore.document import Document as LangDocument
from docx import Document

db = get_fs_client()
PROJECT_ID = get_config("GCP_PROJECT")
GS_BUCKET = get_config("GS_BUCKET")
NAMESPACE = InterviewNamespace.MATERIALS


def index_user_material(user_id, filename, file_id):
    try:
        file_type = identify_file_type(filename)
        if file_type == 'PDF':
            text = pdf_to_text(f"files/{user_id}/{filename}", GS_BUCKET)
        elif file_type == 'DOCX':
            text = docx_to_text(f"files/{user_id}/{filename}", GS_BUCKET)
        elif file_type == 'TXT':
            text = txt_to_text(f"files/{user_id}/{filename}", GS_BUCKET)
        else:
            logging.warning(f"Unsupported file type: {file_type}")
            return json.dumps({"message": f"Unsupported file type: {file_type}"})
        index_ids = UserIndexHelper.index_loader(
            user_id=user_id,
            loader=StringLoader(text),
            namespace=NAMESPACE,
        )
        logging.info(f"Successfully indexed file with index_ids: {index_ids}")

    except Exception as e:
        logging.warning(f"Error indexing file: {str(e)}")
        traceback.print_exc()
        return json.dumps({"message": f"Error indexing file: {str(e)}"})
    document_ref = db.document(f"users/{user_id}/files/{file_id}")
    document_ref.update({
        "index_ids": index_ids,
        "indexed_at": datetime.now(),
    })

    return True


#### Delete document from pinecone
def delete_material_index(user_id, file_id):
    ref = db.document(f"users/{user_id}/files/{file_id}")
    index_ids = []
    try:
        doc = ref.get().to_dict()
        index_ids = doc["index_ids"]
        if index_ids is None:
            return False
        ref.update({"index_ids": fstore.DELETE_FIELD, "is_archived": True})
        index.delete(ids=index_ids, namespace=NAMESPACE.value)
        logging.info(f"Successfully deleted index from user{user_id} file {file_id}")
        return True
    except Exception as e:
        logging.warning(f"Error deleting document: {str(e)}")
        return False


#### Loaders and helper functions
class StringLoader(BaseLoader):
    # load string

    def __init__(
            self,
            text: str,
    ):
        self.text = text

    def load(self) -> List[LangDocument]:
        """Load from text."""
        # Just assign the text directly. There's no need to read from a file.
        text = self.text

        metadata = {"source": "direct input"}  # Change the metadata as the source is not a file
        return [LangDocument(page_content=text, metadata=metadata)]


def identify_file_type(file_name):
    file_type = file_name.split('.')[-1].upper()
    return file_type


def file_loader(file_path, bucket_name):
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(file_path)
    logging.info("<<<Downloading start>>>")
    file_bytes = blob.download_as_bytes()
    logging.info("<<<Downloading end>>>")
    file_io = BytesIO(file_bytes)
    return file_io


def pdf_to_text(file_path, bucket_name):
    logging.info("<<<pdf_to_text start>>>")
    pdf_file_io = file_loader(file_path, bucket_name)

    # Initialize PyMuPDF document from the BytesIO object
    doc = fitz.open(stream=pdf_file_io, filetype="pdf")

    text = []  # List to hold text for each page

    for page in doc:
        # Extract text from each page
        page_text = page.get_text()
        text.append(page_text)  # Add the page text to the list

    logging.info("<<<pdf_to_text end>>>")

    doc.close()  # Close the document

    return '\n'.join(text)


def docx_to_text(file_path, bucket_name):
    logging.info("<<<docx_to_text start>>>")
    docx_file_io = file_loader(file_path, bucket_name)

    # Load DOCX file from BytesIO object
    document = Document(docx_file_io)
    text = [paragraph.text for paragraph in document.paragraphs]

    logging.info("<<<docx_to_text end>>>")
    return '\n'.join(text)


def txt_to_text(file_path, bucket_name):
    logging.info("<<<txt_to_text start>>>")
    txt_file_io = file_loader(file_path, bucket_name)

    # Read text directly from the BytesIO object
    text = txt_file_io.read().decode('utf-8')  # Ensure text is decoded properly

    logging.info("<<<txt_to_text end>>>")
    return text
